"""
generate_samples.py — Dev-only script to create sample resume files (fpdf2 v2.8+).

Generates 8 dummy resume files across PDF, DOCX, and TXT formats
in the 'resumes/' directory for testing the FileSystemAssistant.

Usage:
    python generate_samples.py

Dependencies (in addition to requirements.txt):
    fpdf2, python-docx  (both already in requirements.txt)
"""

import os
from pathlib import Path
from fpdf import FPDF, XPos, YPos
from docx import Document

RESUMES_DIR = Path("resumes")
RESUMES_DIR.mkdir(exist_ok=True)

# ─────────────────────────────────────────────────────────────────────────────
# Resume Data
# ─────────────────────────────────────────────────────────────────────────────

RESUMES = [
    {
        "filename": "john_doe.pdf",
        "format": "pdf",
        "content": """JOHN DOE
Software Engineer | john.doe@email.com | +1-555-0101 | LinkedIn: linkedin.com/in/johndoe

SUMMARY
Experienced software engineer with 6+ years of expertise in backend development.
Passionate about building scalable systems and high-quality APIs.

TECHNICAL SKILLS
- Languages: Python, Go, SQL
- Frameworks: FastAPI, Flask, Django
- Tools: Docker, Kubernetes, PostgreSQL, Redis, Git

EXPERIENCE
Senior Software Engineer — TechCorp Inc.  (2021 – Present)
  • Architected a Python microservices platform serving 2M+ daily requests.
  • Reduced API latency by 40% through caching strategies using Redis.
  • Led a team of 4 engineers for Python backend development initiatives.

Software Engineer — StartupXYZ  (2018 – 2021)
  • Built RESTful APIs using Python and Flask for a SaaS analytics product.
  • Developed ETL pipelines in Python to process 50GB+ daily data volumes.
  • Implemented 5+ years of Python experience building backend applications.

EDUCATION
B.Sc. Computer Science — State University (2018)

CERTIFICATIONS
• AWS Certified Developer – Associate
• Python Institute PCEP Certification
""",
    },
    {
        "filename": "jane_smith.pdf",
        "format": "pdf",
        "content": """JANE SMITH
Data Scientist | jane.smith@email.com | +1-555-0202

SUMMARY
Data scientist with 4 years of experience in machine learning, statistical modelling,
and data analysis. Strong background in Python and deep learning frameworks.

TECHNICAL SKILLS
- Languages: Python, R, SQL
- ML Libraries: TensorFlow, PyTorch, scikit-learn, Pandas, NumPy
- Platforms: AWS SageMaker, Google Colab, Databricks

EXPERIENCE
Data Scientist — DataDriven Corp  (2022 – Present)
  • Developed Python-based ML models achieving 92% accuracy on classification tasks.
  • Built NLP pipelines using Python and HuggingFace Transformers.
  • Conducted A/B testing experiments using Python and statistical analysis.

Junior Data Analyst — Analytics Co.  (2020 – 2022)
  • Created data dashboards using Python (Matplotlib, Seaborn).
  • Wrote complex SQL queries to extract insights from relational databases.

EDUCATION
M.Sc. Data Science — Tech University (2020)
B.Sc. Mathematics — Liberal Arts College (2018)
""",
    },
    {
        "filename": "alice_jones.docx",
        "format": "docx",
        "content": [
            ("ALICE JONES", True),
            ("Project Manager | alice.jones@email.com | +1-555-0303", False),
            ("", False),
            ("SUMMARY", True),
            (
                "Dynamic project manager with 8+ years driving cross-functional teams to deliver "
                "complex software products on time and within budget. Recognized for strong "
                "leadership skills and stakeholder management expertise.",
                False,
            ),
            ("", False),
            ("CORE COMPETENCIES", True),
            ("• Leadership & Team Management", False),
            ("• Agile / Scrum Methodology", False),
            ("• Risk Management & Mitigation", False),
            ("• Budget Planning & Cost Control", False),
            ("• Stakeholder Communication", False),
            ("", False),
            ("EXPERIENCE", True),
            ("Senior Project Manager — Enterprise Solutions Ltd  (2019 – Present)", True),
            (
                "Spearheaded leadership of 12-person cross-functional teams across 3 time zones. "
                "Managed $2.5M project budgets with consistent on-time, on-budget delivery. "
                "Introduced Agile ceremonies that improved sprint velocity by 30%. "
                "Mentored junior PMs, demonstrating strong coaching and leadership skills.",
                False,
            ),
            ("", False),
            ("Project Manager — Digital Ventures  (2016 – 2019)", True),
            (
                "Led delivery of 20+ web and mobile product launches. "
                "Established leadership forums to align engineering and business stakeholders. "
                "Reduced project overruns by 25% through proactive risk management.",
                False,
            ),
            ("", False),
            ("EDUCATION", True),
            ("MBA — Business School of Excellence (2016)", False),
            ("B.Sc. Information Systems — State University (2014)", False),
            ("", False),
            ("CERTIFICATIONS", True),
            ("• PMP — Project Management Professional", False),
            ("• Certified Scrum Master (CSM)", False),
        ],
    },
    {
        "filename": "bob_martin.docx",
        "format": "docx",
        "content": [
            ("BOB MARTIN", True),
            ("Java Backend Engineer | bob.martin@email.com | +1-555-0404", False),
            ("", False),
            ("SUMMARY", True),
            (
                "Java backend engineer with 5 years of experience building enterprise-grade "
                "microservices and distributed systems. Proficient in Spring Boot, Kafka, and "
                "cloud-native architectures.",
                False,
            ),
            ("", False),
            ("TECHNICAL SKILLS", True),
            ("• Languages: Java, Kotlin, SQL", False),
            ("• Frameworks: Spring Boot, Spring Security, Hibernate", False),
            ("• Messaging: Apache Kafka, RabbitMQ", False),
            ("• Cloud: AWS (EC2, S3, RDS), Docker, Kubernetes", False),
            ("", False),
            ("EXPERIENCE", True),
            ("Senior Java Engineer — FinTech Solutions  (2021 – Present)", True),
            (
                "Designed and implemented Java-based microservices processing 500K+ transactions/day. "
                "Built event-driven architectures using Apache Kafka in Java. "
                "Optimized SQL queries reducing database load by 35%.",
                False,
            ),
            ("", False),
            ("Java Developer — Enterprise Apps Inc.  (2019 – 2021)", True),
            (
                "Developed RESTful APIs using Java Spring Boot for enterprise CRM platform. "
                "Implemented OAuth2 security layer using Spring Security. "
                "Wrote unit and integration tests in JUnit achieving 85% code coverage.",
                False,
            ),
            ("", False),
            ("EDUCATION", True),
            ("B.Sc. Computer Science — Technical University (2019)", False),
        ],
    },
    {
        "filename": "carol_white.txt",
        "format": "txt",
        "content": """CAROL WHITE
Database Engineer | carol.white@email.com | +1-555-0505

SUMMARY
Skilled database engineer with 7 years of experience designing, optimizing,
and maintaining large-scale relational and NoSQL database systems.

TECHNICAL SKILLS
- Databases: PostgreSQL, MySQL, Microsoft SQL Server, MongoDB
- Query Languages: SQL, PL/SQL, T-SQL
- Tools: pgAdmin, DBeaver, Liquibase, Apache Airflow
- Cloud: AWS RDS, Azure SQL, Google Cloud Spanner

EXPERIENCE

Senior Database Engineer — DataSystems Corp  (2020 – Present)
  * Architected SQL database schemas supporting 10M+ records for e-commerce platform.
  * Wrote complex SQL stored procedures and triggers improving reporting speed by 50%.
  * Designed disaster recovery strategies achieving 99.99% uptime SLA compliance.
  * Led SQL query optimization sprints reducing avg query time from 8s to under 400ms.

Database Administrator — Legacy Systems Inc.  (2017 – 2020)
  * Managed 30+ SQL Server instances across development, staging, and production.
  * Automated database backups using SQL scripts and Windows Task Scheduler.
  * Migrated 5TB of legacy SQL data to PostgreSQL with zero data loss.

EDUCATION
B.Sc. Information Technology — City College (2017)

CERTIFICATIONS
• Microsoft Certified: Azure Database Administrator Associate
• Oracle Database SQL Certified Associate
""",
    },
    {
        "filename": "dave_brown.txt",
        "format": "txt",
        "content": """DAVE BROWN
Frontend Developer | dave.brown@email.com | +1-555-0606

SUMMARY
Creative frontend developer with 3 years of experience building responsive,
performant web applications using modern JavaScript frameworks.

TECHNICAL SKILLS
- Languages: JavaScript, TypeScript, HTML5, CSS3
- Frameworks: React, Next.js, Vue.js
- Tools: Webpack, Vite, Tailwind CSS, Figma, Git
- Testing: Jest, React Testing Library, Cypress

EXPERIENCE

Frontend Developer — WebAgency Creative  (2023 – Present)
  * Built React component libraries used across 10+ client projects.
  * Implemented Next.js SSR applications improving Core Web Vitals scores by 40%.
  * Collaborated with UX designers translating Figma designs into pixel-perfect React components.
  * Wrote TypeScript interfaces and React hooks improving code maintainability.

Junior Web Developer — StartupFront  (2021 – 2023)
  * Developed responsive web interfaces using React and Tailwind CSS.
  * Integrated REST APIs into React applications using Axios and React Query.
  * Created end-to-end tests with Cypress covering critical user flows.

EDUCATION
B.Sc. Web Technologies — Design & Tech Institute (2021)
""",
    },
    {
        "filename": "emily_chen.txt",
        "format": "txt",
        "content": """EMILY CHEN
Full Stack Developer | emily.chen@email.com | +1-555-0707

SUMMARY
Versatile full stack developer with 5 years building end-to-end web applications.
Comfortable across the entire stack from Python backend APIs to React frontends.

TECHNICAL SKILLS
- Backend: Python, FastAPI, Django REST Framework, Node.js
- Frontend: React, TypeScript, Redux, Tailwind CSS
- Databases: PostgreSQL, MongoDB, Redis
- DevOps: Docker, GitHub Actions, AWS (Lambda, EC2, S3)

EXPERIENCE

Full Stack Developer — ProductLab  (2022 – Present)
  * Built Python FastAPI backends serving 100K+ monthly active users.
  * Developed Python data processing pipelines using Pandas and Celery.
  * Created React dashboards consuming Python REST APIs with real-time WebSocket updates.
  * Containerized Python services using Docker and orchestrated with GitHub Actions CI/CD.

Software Developer — AgencyTech  (2019 – 2022)
  * Developed Python Django applications for 15+ client engagements.
  * Wrote Python unit tests achieving 90% coverage using pytest.
  * Built and maintained React frontends integrated with Python backends.

EDUCATION
B.Sc. Computer Science — Engineering University (2019)

CERTIFICATIONS
• AWS Certified Solutions Architect – Associate
• Python Institute PCAP Certification
""",
    },
    {
        "filename": "michael_ross.txt",
        "format": "txt",
        "content": """MICHAEL ROSS
Machine Learning Engineer | michael.ross@email.com | +1-555-0808

SUMMARY
ML engineer with 4 years of experience productionizing machine learning models
and building ML infrastructure. Expertise in Python-based ML pipelines and MLOps.

TECHNICAL SKILLS
- Languages: Python, Scala
- ML/AI: PyTorch, TensorFlow, scikit-learn, XGBoost, LangChain
- MLOps: MLflow, DVC, Weights & Biases, Kubeflow
- Data Engineering: Apache Spark, Airflow, dbt

EXPERIENCE

Machine Learning Engineer — AI Ventures  (2022 – Present)
  * Designed end-to-end Python ML training pipelines reducing model training time by 60%.
  * Deployed Python-based recommendation models serving 5M+ daily predictions.
  * Built LLM-powered features using Python and LangChain for internal knowledge tools.
  * Implemented ML experiment tracking using Python MLflow achieving full reproducibility.

Data Scientist — AnalyticsHouse  (2020 – 2022)
  * Built Python classification models for churn prediction with 88% F1 score.
  * Created Python feature engineering pipelines processing 500GB weekly.

EDUCATION
M.Sc. Machine Learning — AI Research University (2020)
B.Sc. Statistics — National University (2018)
""",
    },
]


# ─────────────────────────────────────────────────────────────────────────────
# Generator Functions
# ─────────────────────────────────────────────────────────────────────────────


def create_pdf(filepath: Path, content: str) -> None:
    """Generate a simple text-based PDF resume using fpdf2 v2.8+."""
    pdf = FPDF()
    pdf.set_margins(left=20, top=20, right=20)
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    effective_width = pdf.epw  # capture once after add_page
    lines = content.strip().split("\n")
    first_line_done = False

    for line in lines:
        stripped = line.strip()
        # Sanitize to latin-1 (required for Helvetica built-in font)
        stripped = stripped.encode("latin-1", errors="replace").decode("latin-1")

        if not first_line_done and stripped:
            # Candidate name — large bold
            pdf.set_font("Helvetica", style="B", size=16)
            pdf.multi_cell(effective_width, 10, stripped)
            first_line_done = True
        elif stripped == "":
            pdf.ln(3)
        elif stripped.isupper() and 1 < len(stripped) < 40:
            # Section heading
            pdf.ln(2)
            pdf.set_font("Helvetica", style="B", size=11)
            pdf.multi_cell(effective_width, 7, stripped)
        else:
            # Body text
            pdf.set_font("Helvetica", size=10)
            pdf.multi_cell(effective_width, 6, stripped)

    pdf.output(str(filepath))
    print(f"  [OK] Created PDF: {filepath}")


def create_docx(filepath: Path, content: list) -> None:
    """Generate a DOCX resume using python-docx."""
    doc = Document()

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = section.bottom_margin = 914400 // 2  # 0.5 inch
        section.left_margin = section.right_margin = 914400  # 1 inch

    for text, bold in content:
        if text == "":
            doc.add_paragraph("")
        else:
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.bold = bold
            if bold and text.isupper():
                run.font.size = None  # default
            para.paragraph_format.space_after = 0

    doc.save(str(filepath))
    print(f"  [OK] Created DOCX: {filepath}")


def create_txt(filepath: Path, content: str) -> None:
    """Write a plain text resume file."""
    filepath.write_text(content.strip(), encoding="utf-8")
    print(f"  [OK] Created TXT:  {filepath}")


# ─────────────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────────────


def main():
    print(f"\nGenerating sample resumes in '{RESUMES_DIR}/' ...\n")

    for resume in RESUMES:
        filepath = RESUMES_DIR / resume["filename"]
        fmt = resume["format"]
        content = resume["content"]

        if fmt == "pdf":
            create_pdf(filepath, content)
        elif fmt == "docx":
            create_docx(filepath, content)
        elif fmt == "txt":
            create_txt(filepath, content)

    print(f"\nDone! {len(RESUMES)} resume files created in '{RESUMES_DIR}/'.\n")


if __name__ == "__main__":
    main()
